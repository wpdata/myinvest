"""
MyInvest V0.3 - Word Report Generator (T029)
Holdings and risk metrics report using python-docx.
"""

import logging
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Word report generator for holdings and risk analysis.

    Generates professional Word documents with:
    - Holdings table
    - Risk metrics
    - Recommended actions
    """

    def __init__(self):
        """Initialize Word report generator."""
        pass

    def generate_holdings_report(
        self,
        holdings: List[Dict[str, Any]],
        risk_metrics: Dict[str, Any],
        output_path: str,
        title: str = "持仓与风险分析报告"
    ) -> str:
        """Generate Word holdings and risk report.

        Args:
            holdings: List of current holdings
            risk_metrics: Risk metrics dict (VaR, concentration, etc.)
            output_path: Output Word file path
            title: Report title

        Returns:
            str: Path to generated Word file
        """
        logger.info(f"[WordGenerator] Generating report: {output_path}")

        # Create document
        doc = Document()

        # Title
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # Section 1: Holdings Table
        doc.add_heading("一、当前持仓", level=1)

        if holdings:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Headers
            headers = ['股票代码', '数量', '当前价格', '市值', '占比']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                # Bold headers
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add holdings data
            total_value = sum(h.get('market_value', 0) for h in holdings)

            for holding in holdings:
                row_cells = table.add_row().cells
                row_cells[0].text = holding.get('symbol', '')
                row_cells[1].text = str(holding.get('quantity', 0))
                row_cells[2].text = f"¥{holding.get('current_price', 0):.2f}"
                row_cells[3].text = f"¥{holding.get('market_value', 0):,.2f}"

                pct = (holding.get('market_value', 0) / total_value * 100) if total_value > 0 else 0
                row_cells[4].text = f"{pct:.1f}%"

            doc.add_paragraph("")
            doc.add_paragraph(f"总市值: ¥{total_value:,.2f}")
        else:
            doc.add_paragraph("当前无持仓")

        doc.add_paragraph("")

        # Section 2: Risk Metrics
        doc.add_heading("二、风险指标", level=1)

        if risk_metrics:
            # VaR
            var_95 = risk_metrics.get('var_95', 0)
            doc.add_paragraph(f"• 风险价值 (VaR 95%): ¥{var_95:,.2f}")

            # Max Drawdown
            max_dd = risk_metrics.get('max_drawdown_pct', 0)
            doc.add_paragraph(f"• 最大回撤: {max_dd:.2f}%")

            # Sharpe Ratio
            sharpe = risk_metrics.get('sharpe_ratio', 0)
            doc.add_paragraph(f"• 夏普比率: {sharpe:.2f}")

            # Concentration Risk
            concentration = risk_metrics.get('concentration_risk', 'low')
            doc.add_paragraph(f"• 集中度风险: {concentration}")
        else:
            doc.add_paragraph("无风险指标数据")

        doc.add_paragraph("")

        # Section 3: Recommended Actions
        doc.add_heading("三、建议操作", level=1)

        recommendations = self._generate_recommendations(holdings, risk_metrics)

        for rec in recommendations:
            p = doc.add_paragraph(rec, style='List Bullet')

        # Footer: Disclaimer
        doc.add_paragraph("")
        doc.add_paragraph("")
        disclaimer = doc.add_paragraph(
            "免责声明：本报告仅供参考，不构成投资建议。"
            "投资有风险，入市需谨慎。"
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in disclaimer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Save document
        doc.save(output_path)

        logger.info(f"[WordGenerator] Report generated: {output_path}")
        return output_path

    def _generate_recommendations(
        self,
        holdings: List[Dict],
        risk_metrics: Dict
    ) -> List[str]:
        """Generate recommended actions based on holdings and risk."""
        recommendations = []

        # Check concentration
        if holdings:
            total_value = sum(h.get('market_value', 0) for h in holdings)
            for holding in holdings:
                pct = (holding.get('market_value', 0) / total_value * 100) if total_value > 0 else 0
                if pct > 30:
                    recommendations.append(
                        f"注意: {holding.get('symbol')} 占比过高 ({pct:.1f}%)，建议分散风险"
                    )

        # Check drawdown
        max_dd = risk_metrics.get('max_drawdown_pct', 0)
        if max_dd > 20:
            recommendations.append(
                f"警告: 最大回撤 {max_dd:.1f}% 超过20%，建议降低仓位或设置止损"
            )

        # Check Sharpe ratio
        sharpe = risk_metrics.get('sharpe_ratio', 0)
        if sharpe < 1.0:
            recommendations.append(
                f"提示: 夏普比率 {sharpe:.2f} 偏低，策略风险调整后收益不佳"
            )

        if not recommendations:
            recommendations.append("当前风险指标正常，继续执行策略")

        return recommendations
